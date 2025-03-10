import streamlit as st
import tempfile
import os
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from docx.shared import Inches
import base64
from PIL import Image
import io
import time
import json
import magic
import pytesseract
from concurrent.futures import ThreadPoolExecutor
import concurrent.futures
import zipfile
from streamlit import components
import clamd
import pandas as pd
from typing import Generator

# Configuration Management
with open("config.json") as f:
    DEFAULTS = json.load(f)

# Error Codes
ERROR_CODES = {
    1: "Encrypted PDF - Please remove password protection",
    2: "Corrupted File - Try re-saving your PDF",
    3: "Virus Detected - File rejected",
    4: "Invalid File Type - Only PDFs allowed",
}

# Initialize ClamAV
try:
    cd = clamd.ClamdAgnostic()
except:
    st.warning("Virus scanning disabled - ClamAV not available")

# Set page configuration
st.set_page_config(
    page_title="PDF to Word Converter By Umair",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown(f"""
<style>
    /* Previous CSS styles plus accessibility additions */
    .screen-reader-only {{ 
        position: absolute;
        left: -10000px;
        top: auto;
        width: 1px;
        height: 1px;
        overflow: hidden;
    }}
</style>
""", unsafe_allow_html=True)

# Modular Handlers
def image_handler(page, doc, quality: int) -> int:
    image_count = 0
    for img_info in page.get_images(full=True):
        xref = img_info[0]
        base_image = doc.extract_image(xref)
        image_bytes = base_image["image"]
        try:
            img = Image.open(io.BytesIO(image_bytes))
            img = img.convert('RGB')  # Ensure compatibility
            img_stream = io.BytesIO()
            img.save(img_stream, format='JPEG', quality=quality)
            doc.add_picture(img_stream, width=Inches(5))
            image_count += 1
        except Exception as e:
            continue
    return image_count

def table_handler(page) -> list:
    tables = []
    try:
        pdf_page = pdfplumber.open(page).pages[0]
        for table in pdf_page.find_tables():
            tables.append(table.extract())
    except:
        pass
    return tables

# Optimized Processing
def optimized_pdf_processing(pdf_path: str, page_range: str, ocr: bool = False) -> Generator:
    doc = fitz.open(pdf_path)
    total_pages = doc.page_count
    
    # Parse page range
    if page_range.lower() != "all":
        pages = []
        for part in page_range.split(','):
            if '-' in part:
                start, end = map(int, part.split('-'))
                pages.extend(range(start-1, end))
            else:
                pages.append(int(part)-1)
    else:
        pages = range(total_pages)

    with ThreadPoolExecutor() as executor:
        futures = {executor.submit(process_page, doc, pg, ocr): pg for pg in pages}
        for future in concurrent.futures.as_completed(futures):
            yield future.result()

def process_page(doc, page_num: int, ocr: bool) -> dict:
    page_data = {'text': '', 'images': [], 'tables': []}
    try:
        page = doc.load_page(page_num)
        page_data['text'] = page.get_text("text")
        
        if ocr:
            pix = page.get_pixmap()
            img = Image.open(io.BytesIO(pix.tobytes()))
            page_data['text'] += '\n' + pytesseract.image_to_string(img)
            
        page_data['tables'] = table_handler(page)
        return page_data
    except:
        return page_data

# Core Conversion Function
def convert_pdf_to_docx(pdf_file, config: dict):
    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    temp_pdf.write(pdf_file.read())
    temp_pdf.close()
    
    doc = Document()
    conversion_stats = {
        'pages': 0,
        'images': 0,
        'tables': 0,
        'text_blocks': 0
    }

    try:
        with fitz.open(temp_pdf.name) as pdf_doc:
            # Validate PDF
            if pdf_doc.is_encrypted:
                raise Exception(ERROR_CODES[1])
            
            # Process pages
            for page_data in optimized_pdf_processing(temp_pdf.name, config['page_range'], config['ocr']):
                if page_data['text'].strip():
                    doc.add_paragraph(page_data['text'])
                    conversion_stats['text_blocks'] += 1
                
                for table in page_data['tables']:
                    conversion_stats['tables'] += 1
                    word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                    for row_idx, row in enumerate(table):
                        for col_idx, cell in enumerate(row):
                            word_table.cell(row_idx, col_idx).text = str(cell)
                
                conversion_stats['images'] += image_handler(page_data.get('page', None), doc, config['image_quality'])
                conversion_stats['pages'] += 1

        output_path = tempfile.NamedTemporaryFile(delete=False, suffix=config['output_format'])
        doc.save(output_path.name)
        with open(output_path.name, "rb") as f:
            file_data = f.read()
            
        return file_data, conversion_stats
    finally:
        os.unlink(temp_pdf.name)
        os.unlink(output_path.name)

# Enhanced UI Components
def show_tour():
    # Implement guided tour logic
    pass

def virus_scan(file) -> bool:
    try:
        result = cd.instream(file)
        return result['stream'][0] == 'OK'
    except:
        return True  # Bypass if ClamAV unavailable

def main():
    st.markdown('<h1 class="main-header">PDF to Word Converter By Umair</h1>', unsafe_allow_html=True)
    
    # Version Control
    if 'conversion_history' not in st.session_state:
        st.session_state.conversion_history = []

    col1, col2 = st.columns([1, 3])
    
    with col1:
        st.markdown("### 🛠 Settings")
        config = {
            'page_range': st.text_input("Convert pages (e.g., 1-3,5)", "all"),
            'image_quality': st.slider("Image Quality", 1, 100, 85),
            'output_format': st.radio("Output Format", [".docx", ".txt"]),
            'ocr': st.checkbox("Enable OCR (for scanned documents)"),
            'batch': st.checkbox("Batch Processing")
        }

    with col2:
        st.markdown("### 📤 Upload PDF")
        uploaded_files = st.file_uploader("Choose PDF files", type="pdf", 
                                       accept_multiple_files=config['batch'])
        
        if uploaded_files:
            for file in uploaded_files:
                # Input Validation
                file.seek(0)
                mime = magic.from_buffer(file.read(1024), mime=True)
                if mime != 'application/pdf':
                    st.error(ERROR_CODES[4])
                    continue
                
                # Virus Check
                if not virus_scan(file):
                    st.error(ERROR_CODES[3])
                    continue

                # Password Check
                try:
                    with fitz.open(stream=file.read(), filetype="pdf") as test_doc:
                        pass
                except:
                    st.error(ERROR_CODES[1])
                    continue

                # Conversion Logic
                if st.button(f"Convert {file.name}", type="primary"):
                    progress = st.progress(0)
                    try:
                        docx_data, stats = convert_pdf_to_docx(file, config)
                        st.session_state.conversion_history.append({
                            'filename': file.name,
                            'timestamp': time.time(),
                            'stats': stats
                        })

                        # Download Logic
                        st.balloons()
                        filename = file.name.replace(".pdf", f"{config['output_format']}")
                        b64 = base64.b64encode(docx_data).decode()
                        components.html(f"""
                            <script>
                                window.open('data:application/octet-stream;base64,{b64}')
                            </script>
                        """)
                        
                    except Exception as e:
                        st.error(str(e))
                    finally:
                        progress.empty()

if __name__ == "__main__":
    if not os.path.exists("config.json"):
        with open("config.json", "w") as f:
            json.dump({
                "default_quality": 85,
                "default_format": ".docx",
                "max_file_size": 50  # MB
            }, f)
    
    if st.session_state.get('first_visit', True):
        show_tour()
        st.session_state.first_visit = False
    
    main()
